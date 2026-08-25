from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "entrega_tecnica_imssight_20260731_140636" / "documentacion"
DOCX_PATH = OUT_DIR / "MANUAL_EJECUTIVO_INSTALACION_VOLUMETRIA_IMSSIGHT.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 43)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border_bottom(paragraph, color="2E74B5", size="8", space="6"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_para(doc, text="", style=None, bold=False, italic=False, color=INK, size=11, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_kv_table(doc, rows, widths=(2500, 6860), header=None):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, list(widths))
    if header:
        row = table.add_row()
        row.cells[0].merge(row.cells[1])
        cell = row.cells[0]
        set_cell_shading(cell, CALLOUT_FILL)
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)
    for label, value in rows:
        row = table.add_row()
        row.cells[0].text = ""
        row.cells[1].text = ""
        set_cell_shading(row.cells[0], LIGHT_FILL)
        for c in row.cells:
            set_cell_margins(c)
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(label)
        set_run_font(r0, size=10.5, color=DARK_BLUE, bold=True)
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(value)
        set_run_font(r1, size=10.5, color=INK)
    add_para(doc, "", after=4)
    return table


def add_data_table(doc, headers, data, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=10, color=DARK_BLUE, bold=True)
    for row_data in data:
        row = table.add_row()
        for i, value in enumerate(row_data):
            p = row.cells[i].paragraphs[0]
            if i > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_run_font(run, size=9.5, color=INK)
    add_para(doc, "", after=4)
    return table


def add_commands_table(doc, data):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2600, 6760])
    for i, header in enumerate(["Actividad", "Comando"]):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=10, color=DARK_BLUE, bold=True)
    for activity, command in data:
        row = table.add_row()
        for i, value in enumerate([activity, command]):
            p = row.cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(run, size=9.5, color=INK)
    add_para(doc, "", after=4)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p.add_run("\n")
    r2 = p.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    add_para(doc, "", after=2)
    return table


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    r = header.add_run("IMSSight | Manual ejecutivo de instalación y volumetría")
    set_run_font(r, size=9, color=MUTED)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer = section.footer.paragraphs[0]
    footer.text = ""
    r = footer.add_run("Documento administrativo de soporte técnico | Entrega 20260731_140636")
    set_run_font(r, size=9, color=MUTED)

    add_para(doc, "MANUAL EJECUTIVO", bold=True, color=MUTED, size=11, after=4)
    title = add_para(
        doc,
        "Instalación, volumetría y puesta en operación de IMSSight",
        bold=True,
        color=RGBColor(0, 0, 0),
        size=22,
        after=4,
    )
    subtitle = add_para(
        doc,
        "Documento de referencia para coordinación de alto nivel y ejecución técnica controlada",
        color=MUTED,
        size=13,
        after=12,
    )
    meta = [
        ("Fecha de corte", "31 de julio de 2026"),
        ("Paquete de entrega", "entrega_tecnica_imssight_20260731_140636"),
        ("Alcance operativo", "Instalación del contenido incluido en la memoria Kingston, sin carga de videos locales."),
        ("Base de datos", "MariaDB 10.11.17, base imssight, respaldo SQL incluido."),
        ("Aplicación", "PHP 8.2.31 sobre Apache 2.4.67, ejecutable mediante Docker Compose."),
        ("Carácter del documento", "Manual ejecutivo-administrativo con comandos mínimos de instalación y verificación."),
    ]
    add_kv_table(doc, meta, header="Ficha de control documental")
    rule = add_para(doc, "", after=8)
    set_paragraph_border_bottom(rule)

    add_heading(doc, "1. Propósito y alcance", 1)
    add_para(
        doc,
        "El presente manual establece, en lenguaje administrativo y ejecutivo, los elementos mínimos para instalar, validar y resguardar la plataforma IMSSight a partir de la memoria de entrega. Su énfasis principal es la volumetría del paquete, las versiones de software observadas, los puntos de montaje y los comandos de ejecución indispensables para el personal técnico responsable.",
    )
    add_callout(
        doc,
        "Criterio rector",
        "La entrega no incorpora videos locales. La carpeta app/video fue excluida del paquete operativo y no debe cargarse al servidor aplicativo. Esta decisión reduce sustancialmente el volumen de transferencia, almacenamiento y ancho de banda requerido.",
    )

    doc.add_page_break()
    add_heading(doc, "2. Contenido visible en la memoria", 1)
    add_data_table(
        doc,
        ["Elemento", "Archivo o ruta", "Volumen aproximado", "Finalidad"],
        [
            ("Código aplicativo", "codigo/imssight_codigo_sin_videos_20260731_140636.tar.gz", "212 MB", "Fuente operativo sin videos locales."),
            ("Base SQL", "database/imssight_20260731_140636.sql", "56 MB", "Restauración completa de la base imssight."),
            ("Base SQL comprimida", "database/imssight_20260731_140636.sql.gz", "3.4 MB", "Respaldo alterno para traslado eficiente."),
            ("Guía técnica", "documentacion/GUIA_TECNICA_DESPLIEGUE_IMSSIGHT.md", "17 KB", "Referencia técnica ampliada."),
            ("Integridad", "CHECKSUMS_SHA256.txt", "334 B", "Validación criptográfica de archivos principales."),
            ("Lectura inicial", "README_ENTREGA.md", "887 B", "Orientación rápida de contenido."),
        ],
        [1850, 3400, 1450, 2660],
    )

    add_heading(doc, "3. Volumetría ejecutiva", 1)
    add_para(
        doc,
        "La volumetría de esta entrega es administrativamente manejable para traslado físico, resguardo y transferencia en red institucional. La reducción principal proviene de excluir videos locales, que en el repositorio original representaban aproximadamente 6.5 GB.",
    )
    add_data_table(
        doc,
        ["Concepto", "Volumen", "Lectura administrativa"],
        [
            ("Carpeta final en Kingston", "272 MB aprox.", "Tamaño total para custodia y traslado."),
            ("Paquete de código sin videos", "212 MB", "Unidad principal de instalación aplicativo."),
            ("SQL sin comprimir", "56 MB", "Restauración transparente y auditable."),
            ("SQL comprimido", "3.4 MB", "Copia práctica para respaldos o transferencia."),
            ("Base lógica reportada por MariaDB", "72 MB", "Tamaño interno aproximado de datos e índices."),
            ("Videos excluidos", "6.5 GB", "No forman parte del despliegue ni deben copiarse al servidor."),
        ],
        [3000, 1800, 4560],
    )
    add_data_table(
        doc,
        ["Objeto de base", "Filas estimadas", "Tamaño aproximado"],
        [
            ("demanda_unidad_mensual_geo", "281,559", "46.58 MB"),
            ("prediccion_demanda", "80,114", "21.09 MB"),
            ("escenas", "140", "1.53 MB"),
            ("search_index", "259", "1.53 MB"),
            ("Objetos totales en base", "43", "72 MB lógicos"),
        ],
        [4200, 2400, 2760],
    )
    add_callout(
        doc,
        "Implicación de capacidad",
        "Sin videos locales, el factor crítico deja de ser almacenamiento bruto y pasa a ser la eficiencia de imágenes, documentos, consultas de base de datos y concurrencia de usuarios. Para 100,000 usuarios no simultáneos, el ancho de banda dependerá del patrón real de navegación y del aprovechamiento de caché.",
    )

    add_heading(doc, "4. Versiones de software de referencia", 1)
    add_data_table(
        doc,
        ["Componente", "Versión observada", "Uso dentro de la entrega"],
        [
            ("PHP", "8.2.31", "Ejecución del backend aplicativo."),
            ("Apache", "2.4.67 Debian", "Servidor web dentro del contenedor."),
            ("MariaDB", "10.11.17-MariaDB-ubu2204", "Motor de base de datos."),
            ("Imagen MariaDB", "mariadb:10.11", "Contenedor de base."),
            ("Imagen web", "php:8.2-apache", "Base del contenedor aplicativo."),
            ("Docker Engine", "29.5.2", "Motor de contenedores observado durante la preparación."),
            ("Docker Compose", "v5.1.3", "Orquestación local del servicio."),
        ],
        [2500, 2600, 4260],
    )

    add_heading(doc, "5. Puntos de montaje y puertos", 1)
    add_data_table(
        doc,
        ["Elemento", "Punto o valor", "Descripción"],
        [
            ("Memoria de entrega", "/Volumes/KINGSTON/entrega_tecnica_imssight_20260731_140636", "Ubicación física observada en macOS."),
            ("Código en contenedor", "./app:/var/www/html", "Montaje del aplicativo PHP/HTML/JS/CSS."),
            ("Datos MariaDB", "mariadb_data:/var/lib/mysql", "Volumen persistente de base de datos."),
            ("Puerto web", "8080:80", "Exposición HTTP del aplicativo."),
            ("Puerto base", "3307:3306", "Exposición local de MariaDB."),
            ("phpMyAdmin", "8081:80", "Administración auxiliar; no exponer públicamente."),
        ],
        [2700, 3300, 3360],
    )

    doc.add_page_break()
    add_heading(doc, "6. Comandos mínimos de instalación", 1)
    add_para(doc, "Los comandos siguientes resumen la operación esperada a partir del contenido de la memoria. Deben ejecutarse por personal técnico con permisos suficientes en el servidor de destino.")
    commands = [
        ("Preparar ruta", "mkdir -p /opt/imssight && cd /opt/imssight"),
        ("Copiar entrega desde memoria", "cp -R /ruta/KINGSTON/entrega_tecnica_imssight_20260731_140636 ./"),
        ("Extraer código", "tar -xzf entrega_tecnica_imssight_20260731_140636/codigo/\nimssight_codigo_sin_videos_20260731_140636.tar.gz"),
        ("Entrar al proyecto", "cd imssight"),
        ("Levantar servicios", "docker compose up -d --build"),
        ("Restaurar base", "docker exec -i imssight-db mariadb -uroot -p < ../entrega_tecnica_imssight_20260731_140636/database/\nimssight_20260731_140636.sql"),
        ("Validar contenedores", "docker ps"),
        ("Validar respuesta web", "curl -I http://localhost:8080/"),
    ]
    add_commands_table(doc, commands)

    add_heading(doc, "7. Verificación posterior", 1)
    for item in [
        "Confirmar que el paquete instalado no contiene la ruta app/video.",
        "Confirmar que el contenedor web se encuentre activo y responda en el puerto 8080.",
        "Confirmar que MariaDB se encuentre activo y conserve el volumen mariadb_data.",
        "Confirmar que la base imssight fue importada con 43 objetos.",
        "Validar acceso a pantalla principal, inicio de sesión, búsqueda, muro, chat, exámenes y documentos.",
        "Validar los checksums incluidos antes de una entrega formal o resguardo institucional.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. Consideraciones de ancho de banda", 1)
    add_para(
        doc,
        "Para una población de 100,000 usuarios no simultáneos, el comportamiento esperado dependerá de sesiones reales, frecuencia de acceso y cacheo de recursos estáticos. Al no incluir videos, una sesión típica puede mantenerse en un rango administrable, sujeto a optimización de imágenes y reutilización de caché del navegador.",
    )
    add_data_table(
        doc,
        ["Supuesto mensual", "Estimación", "Comentario ejecutivo"],
        [
            ("100,000 sesiones x 5 MB", "0.5 TB/mes", "Escenario optimizado con caché e imágenes contenidas."),
            ("100,000 sesiones x 15 MB", "1.5 TB/mes", "Escenario intermedio razonable para plataforma educativa sin video local."),
            ("100,000 sesiones x 25 MB", "2.5 TB/mes", "Escenario conservador por navegación amplia e imágenes pesadas."),
        ],
        [3000, 2100, 4260],
    )
    add_callout(
        doc,
        "Recomendación ejecutiva",
        "Mantener fuera del servidor cualquier video futuro. Para preservar estabilidad institucional, los recursos pesados deben administrarse como contenido externo o mediante infraestructura de distribución especializada.",
    )

    add_heading(doc, "9. Criterios de aceptación", 1)
    for item in [
        "La memoria contiene el paquete de código sin videos, la base SQL y la documentación asociada.",
        "El despliegue puede ejecutarse con Docker Compose a partir del contenido incluido.",
        "La base de datos se restaura sin errores y conserva los objetos reportados.",
        "El aplicativo responde por HTTP y las funciones principales son verificables.",
        "La volumetría entregada es congruente con una distribución sin videos locales.",
    ]:
        add_number(doc, item)

    add_heading(doc, "10. Cierre administrativo", 1)
    add_para(
        doc,
        "La presente entrega debe considerarse apta para transferencia técnica controlada, revisión de coordinación y ejecución por personal de infraestructura. Su alcance queda delimitado al contenido visible en la memoria y excluye expresamente la carga de videos locales al servidor aplicativo.",
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())
